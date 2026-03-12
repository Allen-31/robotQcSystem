# -*- coding: utf-8 -*-
"""按规范模板生成接口文档 Word（含：接口描述、URL、请求方式、请求字段、请求示例、返回字段、返回示例）"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.add_run('').font.size = Pt(16)
    return p

def add_table_with_headers(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            if col_idx < len(table.rows[row_idx + 1].cells):
                table.rows[row_idx + 1].cells[col_idx].text = str(cell_text)
    doc.add_paragraph()

def add_json_example(doc, label, json_str):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.paragraph_format.space_after = Pt(3)
    p2 = doc.add_paragraph(json_str)
    p2.paragraph_format.left_indent = Cm(0.5)
    for run in p2.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph()

def add_interface(doc, title, desc, url, method, req_fields, req_example, res_fields, res_example):
    """添加一个完整接口（7 部分）"""
    add_heading(doc, title, level=1)
    # 1. 接口描述
    doc.add_paragraph('1. 接口描述')
    doc.add_paragraph(desc)
    # 2. URL
    doc.add_paragraph('2. URL')
    doc.add_paragraph(url)
    # 3. 请求方式
    doc.add_paragraph('3. 请求方式')
    doc.add_paragraph(method)
    # 4. 请求字段列表
    doc.add_paragraph('4. 请求字段')
    if req_fields:
        add_table_with_headers(doc, ['参数名', '必填', '类型', '说明'], req_fields)
    else:
        doc.add_paragraph('无请求体（GET 时可为查询参数，见 URL 说明）。')
    # 5. 请求参数代码示意
    doc.add_paragraph('5. 请求参数示例')
    add_json_example(doc, '', req_example if req_example else '{}')
    # 6. 返回字段列表
    doc.add_paragraph('6. 返回字段')
    add_table_with_headers(doc, ['参数名', '必填', '类型', '说明'], res_fields)
    # 7. 返回字段代码示意
    doc.add_paragraph('7. 返回示例')
    add_json_example(doc, '', res_example)
    doc.add_paragraph()
    doc.add_paragraph('—' * 30)

def main():
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)

    # 标题
    title = doc.add_paragraph()
    title.add_run('云平台后端接口需求说明（按前端功能反推）').bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    doc.add_paragraph('本文档按《云平台接口规范》标准格式整理，包含：接口描述、URL、请求方式、请求字段列表、请求参数示例、返回字段列表、返回示例。')
    doc.add_paragraph('统一响应外壳：{ "code": 200, "msg": "success", "data": { ... }, "timestamp": 1672531200, "traceId": "uuid4" }，以下返回示例中 data 内为业务字段。')
    doc.add_paragraph()

    # ========== 一、已联调接口 ==========
    add_heading(doc, '一、已联调接口（保持不变）', level=1)
    doc.add_paragraph('以下为当前前后端已联调接口，保持现有路径与字段不变。')
    doc.add_paragraph()

    # 1.1 登录
    add_interface(
        doc,
        '1.1 登录',
        '前端登录页提交用户名、密码后调用此接口，验证通过后返回 token 与用户信息，用于后续请求鉴权。',
        '/api/auth/login',
        'POST',
        [
            ('username', 'Y', 'string', '用户名（登录账号）'),
            ('password', 'Y', 'string', '密码'),
            ('remember', 'N', 'boolean', '是否记住登录状态'),
        ],
        '''{
    "username": "admin",
    "password": "123456",
    "remember": true
}''',
        [
            ('token', 'Y', 'string', '访问令牌，请求头携带'),
            ('user', 'Y', 'object', '用户信息'),
            ('user.code', 'Y', 'string', '用户编码'),
            ('user.displayName', 'Y', 'string', '显示名称'),
            ('user.roles', 'Y', 'array', '角色编码列表'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": {
        "token": "eyJhbGc...",
        "user": {
            "code": "admin",
            "displayName": "系统管理员",
            "roles": ["admin"]
        }
    },
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.2 登出
    add_interface(
        doc,
        '1.2 登出',
        '用户登出时调用，服务端使当前 token 失效；请求需携带 token。',
        '/api/auth/logout',
        'POST',
        [],
        '{}',
        [
            ('（无 body 或 success）', 'N', '-', '按现有后端约定'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": {},
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.3 获取当前用户
    add_interface(
        doc,
        '1.3 获取当前用户信息',
        '根据当前 token 获取当前登录用户信息，用于刷新用户及权限。',
        '/api/auth/me',
        'GET',
        [],
        '（无请求体）',
        [
            ('code', 'Y', 'string', '用户编码'),
            ('displayName', 'Y', 'string', '显示名称'),
            ('roles', 'Y', 'array', '角色编码列表'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": {
        "code": "admin",
        "displayName": "系统管理员",
        "roles": ["admin"]
    },
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.4 用户分页列表
    add_interface(
        doc,
        '1.4 用户管理 - 分页列表',
        '部署配置-用户管理页获取用户分页列表，支持关键词、角色、状态筛选。',
        '/api/deploy/users?page=1&pageSize=20&keyword=xxx&role=xxx&status=xxx',
        'GET',
        [
            ('page', 'N', 'number', '页码，默认 1'),
            ('pageSize', 'N', 'number', '每页条数，默认 20'),
            ('keyword', 'N', 'string', '关键词（编码/姓名等）'),
            ('role', 'N', 'string', '角色编码筛选'),
            ('status', 'N', 'string', '状态：enabled/disabled'),
        ],
        '（GET 无 body，参数在 URL 查询）',
        [
            ('list', 'Y', 'array', '用户列表'),
            ('list[].code', 'Y', 'string', '用户编码'),
            ('list[].name', 'Y', 'string', '姓名'),
            ('list[].phone', 'N', 'string', '手机号'),
            ('list[].email', 'N', 'string', '邮箱'),
            ('list[].status', 'Y', 'string', '状态'),
            ('list[].lastLoginAt', 'N', 'string', '最后登录时间'),
            ('list[].roles', 'Y', 'array', '角色编码列表'),
            ('total', 'Y', 'number', '总条数'),
            ('pageNum', 'N', 'number', '当前页'),
            ('pageSize', 'N', 'number', '每页条数'),
            ('pages', 'N', 'number', '总页数'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": {
        "list": [
            {
                "code": "admin",
                "name": "管理员",
                "phone": null,
                "email": null,
                "status": "enabled",
                "lastLoginAt": "2026-03-06 10:00:00",
                "roles": ["admin"]
            }
        ],
        "total": 1,
        "pageNum": 1,
        "pageSize": 20,
        "pages": 1
    },
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.5 新增用户
    add_interface(
        doc,
        '1.5 用户管理 - 新增用户',
        '部署配置-用户管理页新建用户。',
        '/api/deploy/users',
        'POST',
        [
            ('code', 'Y', 'string', '用户编码'),
            ('name', 'Y', 'string', '姓名'),
            ('phone', 'N', 'string', '手机号'),
            ('email', 'N', 'string', '邮箱'),
            ('status', 'N', 'string', '状态：enabled/disabled'),
            ('roles', 'N', 'array', '角色编码列表'),
            ('password', 'Y', 'string', '登录密码'),
        ],
        '''{
    "code": "newuser",
    "name": "新用户",
    "phone": "13800138000",
    "email": "user@example.com",
    "status": "enabled",
    "roles": ["qc"],
    "password": "123456"
}''',
        [
            ('data', 'N', 'null/object', '按现有后端约定'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.6 更新用户
    add_interface(
        doc,
        '1.6 用户管理 - 更新用户',
        '部署配置-用户管理页编辑用户信息（不包含密码）。',
        '/api/deploy/users/{code}',
        'PUT',
        [
            ('name', 'Y', 'string', '姓名'),
            ('phone', 'N', 'string', '手机号'),
            ('email', 'N', 'string', '邮箱'),
            ('status', 'N', 'string', '状态'),
            ('roles', 'N', 'array', '角色编码列表'),
        ],
        '''{
    "name": "新用户",
    "phone": "13800138000",
    "email": "user@example.com",
    "status": "enabled",
    "roles": ["qc"]
}''',
        [('data', 'N', '-', '按现有约定')],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.7 删除用户
    add_interface(
        doc,
        '1.7 用户管理 - 删除用户',
        '部署配置-用户管理页删除指定用户。',
        '/api/deploy/users/{code}',
        'DELETE',
        [],
        '（无请求体，code 在路径中）',
        [('data', 'N', '-', '按现有约定')],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.8 更新用户角色
    add_interface(
        doc,
        '1.8 用户管理 - 更新用户角色',
        '部署配置-用户管理页为指定用户设置所属角色。',
        '/api/deploy/users/{code}/roles',
        'PUT',
        [('roles', 'Y', 'array', '角色编码列表')],
        '''{
    "roles": ["admin", "qc"]
}''',
        [('data', 'N', '-', '按现有约定')],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.9 修改密码
    add_interface(
        doc,
        '1.9 用户管理 - 修改密码',
        '部署配置-用户管理页为指定用户修改密码（需原密码）。',
        '/api/deploy/users/{code}/password',
        'PUT',
        [
            ('oldPassword', 'Y', 'string', '原密码'),
            ('newPassword', 'Y', 'string', '新密码'),
        ],
        '''{
    "oldPassword": "123456",
    "newPassword": "654321"
}''',
        [
            ('success', 'N', 'boolean', '是否成功'),
            ('error', 'N', 'string', '错误码，如 old_password_invalid'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": { "success": true },
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.10 角色列表
    add_interface(
        doc,
        '1.10 角色管理 - 角色列表',
        '部署配置-角色管理页获取角色列表，支持关键词筛选。',
        '/api/deploy/roles?keyword=xxx',
        'GET',
        [('keyword', 'N', 'string', '关键词（编码/名称/描述）')],
        '（GET 无 body）',
        [
            ('（data 为数组）', 'Y', 'array', '角色列表'),
            ('[].code', 'Y', 'string', '角色编码'),
            ('[].name', 'Y', 'string', '角色名称'),
            ('[].description', 'Y', 'string', '描述'),
            ('[].memberCount', 'N', 'number', '成员数'),
            ('[].updatedAt', 'N', 'string', '更新时间'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": [
        {
            "code": "admin",
            "name": "超级管理员",
            "description": "系统默认角色",
            "memberCount": 1,
            "updatedAt": "2026-03-06 10:00:00"
        }
    ],
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.11 新增角色
    add_interface(
        doc,
        '1.11 角色管理 - 新增角色',
        '部署配置-角色管理页新建角色。',
        '/api/deploy/roles',
        'POST',
        [
            ('code', 'Y', 'string', '角色编码'),
            ('name', 'Y', 'string', '角色名称'),
            ('description', 'N', 'string', '描述'),
        ],
        '''{
    "code": "operator",
    "name": "操作员",
    "description": "普通操作员角色"
}''',
        [('data', 'N', '-', '按现有约定')],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.12 更新角色
    add_interface(
        doc,
        '1.12 角色管理 - 更新角色',
        '部署配置-角色管理页编辑角色信息。',
        '/api/deploy/roles/{code}',
        'PUT',
        [
            ('name', 'Y', 'string', '角色名称'),
            ('description', 'N', 'string', '描述'),
        ],
        '''{
    "name": "操作员",
    "description": "普通操作员角色"
}''',
        [('data', 'N', '-', '按现有约定')],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.13 删除角色
    add_interface(
        doc,
        '1.13 角色管理 - 删除角色',
        '部署配置-角色管理页删除指定角色（至少保留一个角色等约束由后端校验）。',
        '/api/deploy/roles/{code}',
        'DELETE',
        [],
        '（无请求体）',
        [('data', 'N', '-', '按现有约定')],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.14 获取角色权限
    add_interface(
        doc,
        '1.14 权限管理 - 获取角色权限',
        '分配权限弹窗加载指定角色的菜单与按钮权限配置。',
        '/api/deploy/roles/{code}/permissions',
        'GET',
        [],
        '（无请求体）',
        [
            ('（data 为数组）', 'Y', 'array', '权限项列表'),
            ('[].menuKey', 'Y', 'string', '菜单/节点 key（与前端菜单树一致）'),
            ('[].actions', 'Y', 'array', '该节点下的动作列表，如 ["display","create","edit"]'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": [
        { "menuKey": "/qualityInspection", "actions": ["display"] },
        { "menuKey": "/qualityInspection/workstationManage", "actions": ["display", "create", "edit", "delete"] }
    ],
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 1.15 保存角色权限
    add_interface(
        doc,
        '1.15 权限管理 - 保存角色权限',
        '分配权限弹窗保存指定角色的菜单与按钮权限配置。',
        '/api/deploy/roles/{code}/permissions',
        'PUT',
        [
            ('permissions', 'Y', 'array', '权限项列表'),
            ('permissions[].menuKey', 'Y', 'string', '菜单 key'),
            ('permissions[].actions', 'Y', 'array', '动作列表'),
        ],
        '''{
    "permissions": [
        { "menuKey": "/qualityInspection", "actions": ["display"] },
        { "menuKey": "/qualityInspection/workstationManage", "actions": ["display", "create", "edit", "delete"] }
    ]
}''',
        [('data', 'N', '-', '按现有约定')],
        '''{
    "code": 200,
    "msg": "success",
    "data": null,
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # ========== 二、待实现接口（示例：下发任务 + 地图列表） ==========
    add_heading(doc, '二、待实现接口（按前端功能反推，示例格式）', level=1)
    doc.add_paragraph('以下按《云平台接口规范》示例格式编写 2 个代表接口，其余模块可参照同一模板补充。')
    doc.add_paragraph()

    # 2.1 下发任务（规范文档示例）
    add_interface(
        doc,
        '2.1 下发任务接口',
        '外部系统可使用此接口，下发任务至调度系统。',
        '/api/task/create',
        'POST',
        [
            ('externalTaskId', 'Y', 'string', '外部系统的任务唯一ID，该任务完成后回传上游系统'),
            ('taskFlowId', 'Y', 'string', '任务流程唯一ID，CPS 收到任务后根据此 ID 找到任务流程并创建实例'),
            ('priority', 'N', 'uint', '任务执行优先级，值越大越优先；空时默认 0；范围 [0,999]'),
            ('parameters', 'N', 'json', '任务动态参数，创建模板时传入'),
        ],
        '''{
    "externalTaskId": "uuid",
    "taskFlowId": "uuid",
    "priority": 1,
    "parameters": "{\\"robotCode\\":\\"robot001\\"}"
}''',
        [
            ('externalTaskId', 'Y', 'string', '外部系统的任务唯一ID'),
            ('taskId', 'Y', 'string', 'CPS 系统的唯一任务ID'),
            ('taskStatus', 'Y', 'string', '任务执行状态：PENDING/RUNNING/SUCCESS/FAILED/CANCELLED'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": {
        "externalTaskId": "uuid",
        "taskId": "uuid",
        "taskStatus": "RUNNING"
    },
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    # 2.2 地图列表（待实现示例）
    add_interface(
        doc,
        '2.2 地图管理 - 分页列表',
        '部署配置-地图管理页获取地图分页列表，支持关键词筛选。',
        '/api/maps?page=1&pageSize=20&keyword=xxx',
        'GET',
        [
            ('page', 'N', 'number', '页码'),
            ('pageSize', 'N', 'number', '每页条数'),
            ('keyword', 'N', 'string', '关键词'),
        ],
        '（GET 无 body）',
        [
            ('list', 'Y', 'array', '地图列表'),
            ('list[].mapCode', 'Y', 'string', '地图编码'),
            ('list[].name', 'Y', 'string', '名称'),
            ('list[].type', 'N', 'string', '类型'),
            ('list[].editStatus', 'N', 'string', '编辑状态'),
            ('list[].publishStatus', 'N', 'string', '发布状态'),
            ('list[].editedAt', 'N', 'string', '编辑时间'),
            ('list[].publishedAt', 'N', 'string', '发布时间'),
            ('total', 'Y', 'number', '总条数'),
        ],
        '''{
    "code": 200,
    "msg": "success",
    "data": {
        "list": [
            {
                "mapCode": "MAP001",
                "name": "车间地图",
                "type": "workshop",
                "editStatus": "completed",
                "publishStatus": "published",
                "editedAt": "2026-03-06 10:00:00",
                "publishedAt": "2026-03-06 11:00:00"
            }
        ],
        "total": 1
    },
    "timestamp": 1672531200,
    "traceId": "uuid4"
}''',
    )

    doc.add_paragraph()
    add_heading(doc, '三、其余待实现接口清单（可依上例展开）', level=1)
    doc.add_paragraph('以下仅列 URL 与方法，请求/返回字段可按上述标准格式补充：')
    doc.add_paragraph()
    others = [
        ('设置', 'GET/PUT', '/api/settings 或 /api/deploy/settings'),
        ('地图新增/更新/删除', 'POST/PUT/DELETE', '/api/maps、/api/maps/{code}'),
        ('设备管理', 'GET/POST/PUT/DELETE', '/api/devices、/api/devices/{code}'),
        ('任务模板', 'GET/POST/PUT/DELETE', '/api/task-templates、/api/task-templates/{code}'),
        ('动作模板', 'GET/POST/PUT/DELETE', '/api/action-templates、/api/action-templates/{code}'),
        ('充电策略', 'GET/POST/PUT/DELETE', '/api/charge-strategies、/api/charge-strategies/{code}'),
        ('归巢策略', 'GET/POST/PUT/DELETE', '/api/homing-strategies、/api/homing-strategies/{code}'),
        ('文件管理', 'GET/POST/DELETE', '/api/files、/api/files/upload、/api/files/{id}'),
        ('机器人类型/零部件/分组', 'GET/POST/PUT/DELETE', '/api/robot-types、/api/robot-parts、/api/robot-groups'),
        ('任务流设计', 'GET/POST/PUT', '/api/task-flows、/api/task-flows/{id}'),
        ('发布管理', 'GET/POST/PUT', '/api/publishes、/api/publishes/{id}/cancel'),
        ('任务管理', 'GET/PUT', '/api/tasks、/api/tasks/{id}/cancel'),
        ('机器人管理', 'GET', '/api/robots、/api/robots/{code}'),
        ('工单管理', 'GET/POST/PUT/DELETE', '/api/work-orders、/api/work-orders/{id}/review、/cancel'),
        ('复检记录', 'GET', '/api/reinspection-records'),
        ('质检区/站配置', 'GET/POST/PUT/DELETE', '/api/inspection-areas、/api/inspection-stations 等'),
        ('运维日志', 'GET', '/api/login-logs、/api/operation-logs、/api/api-logs、/api/exception-notifications'),
        ('数据统计', 'GET', '/api/quality-statistics、/api/quality-reports、/api/device-statistics、/api/exception-statistics'),
        ('运营监控/首页', 'GET', '/api/operation-monitoring、/api/dashboard'),
    ]
    for name, method, url in others:
        doc.add_paragraph(f'• {name}：{method}  {url}', style='List Bullet')

    out_path = 'docs/后端接口需求说明-按前端功能反推.docx'
    doc.save(out_path)
    print(f'已生成：{out_path}')

if __name__ == '__main__':
    main()
